import json
import os
import re
import shutil
import time
import xml.etree.ElementTree as ET
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any, Dict, List, Tuple

from bs4 import BeautifulSoup
from docling.document_converter import DocumentConverter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from file_service import FileService
from openai import OpenAI, OpenAIError
from tenacity import retry, stop_after_attempt, wait_exponential
from unstructured.partition.auto import partition

from .prompts import TABLE_SUMMARY_PROMPT, TABLE_SYSTEM_MESSAGE, VISION_API_PROMPT


class DocumentPreProcessor:
    """
    Cleans and preprocesses documents for further processing in the pipeline.

    This class supports flexible integration with external AI providers (e.g., OpenAI) for advanced document enrichment, such as table and image summarization. API credentials and provider type can be specified via parameters for extensibility.

    Args:
        data_dir (str): Path to the central data directory (containing input, output, images, temp subfolders).
        steps (int, optional): Number of cleaning steps to perform.
        remove_headers (bool, optional): Remove headers from the document.
        remove_footers (bool, optional): Remove footers from the document.
        remove_toc (bool, optional): Remove table of contents.
        remove_empty (bool, optional): Remove empty paragraphs or sections.
        abbreviation_strategy (str, optional): Strategy for handling abbreviations.
        footnote_handling (str, optional): Strategy for handling footnotes.
        process_table_with_ai (bool, optional): Use AI to process tables.
        process_images_with_ai (bool, optional): Use AI to process images.
        api_key (str, optional): API key for external AI provider (e.g., OpenAI).
        api_provider (str, optional): Name of the external AI provider (e.g., "openai").
    """

    def __init__(
        self,
        data_dir: str,
        steps: int = 6,
        remove_headers: bool = True,
        remove_footers: bool = True,
        remove_toc: bool = True,
        remove_empty: bool = True,
        abbreviation_strategy: str = "inline",
        footnote_handling: str = "remove",
        process_table_with_ai: bool = True,
        process_images_with_ai: bool = True,
        api_key: str = None,
        api_provider: str = None,
    ) -> None:
        self.data_dir = data_dir
        self.input_dir = str(Path(data_dir) / "input")
        self.output_dir = str(Path(data_dir) / "output")
        self.images_dir = str(Path(data_dir) / "images")
        self.temp_dir = str(Path(data_dir) / "temp")
        self.steps = steps
        self.remove_headers = remove_headers
        self.remove_footers = remove_footers
        self.remove_toc = remove_toc
        self.remove_empty = remove_empty
        self.abbreviation_strategy = abbreviation_strategy
        self.footnote_handling = footnote_handling
        self.process_table_with_ai = process_table_with_ai
        self.process_images_with_ai = process_images_with_ai
        self.api_key = api_key
        self.api_provider = api_provider

    def run(self, input_filename: str) -> None:
        """
        Run the document preprocessing pipeline for a given input and output filename (relative to input_dir/output_dir).
        The input file is copied to the temp directory and all processing is done on the temp copy to protect the original.
        """

        client = None
        if self.api_key and self.api_provider:
            if self.api_provider == "openai":
                client = OpenAI(api_key=self.api_key)
                self.text_model = "gpt-4o-mini"
                self.image_model = "gpt-4o-mini"
        else:
            self.text_model = None
            self.image_model = None

        input_path = str(Path(self.input_dir) / input_filename)
        output_path = str(Path(self.output_dir) / input_filename)

        temp_copy_path = str(Path(self.temp_dir) / f"working_{input_filename}")
        shutil.copy2(input_path, temp_copy_path)

        self.process_data(
            client=client,
            input_path=temp_copy_path,
            output_path=output_path,
            steps=self.steps,
            remove_headers=self.remove_headers,
            remove_footers=self.remove_footers,
            remove_toc=self.remove_toc,
            remove_empty=self.remove_empty,
            abbreviation_strategy=self.abbreviation_strategy,
            footnote_handling=self.footnote_handling,
            process_table_with_ai=self.process_table_with_ai,
            process_images_with_ai=self.process_images_with_ai,
        )

    def process_data(
        self,
        client: object = None,
        input_path: str = None,
        output_path: str = None,
        steps: int = 6,
        remove_headers: bool = True,
        remove_footers: bool = True,
        remove_toc: bool = True,
        remove_empty: bool = True,
        abbreviation_strategy: str = "inline",
        footnote_handling: str = "remove",
        process_table_with_ai: bool = True,
        process_images_with_ai: bool = True,
    ) -> None:
        """
        Process the document with the specified cleaning and enrichment steps.

        Args:
            client (object, optional): External AI client (e.g., OpenAI). If None, AI steps are skipped.
            input_path (str, optional): Path to the input document.
            output_path (str, optional): Path to the output document.
            steps (int, optional): Number of cleaning steps to perform.
            remove_headers (bool, optional): Remove headers from the document.
            remove_footers (bool, optional): Remove footers from the document.
            remove_toc (bool, optional): Remove table of contents.
            remove_empty (bool, optional): Remove empty paragraphs or sections.
            abbreviation_strategy (str, optional): Strategy for handling abbreviations.
            footnote_handling (str, optional): Strategy for handling footnotes.
            process_table_with_ai (bool, optional): Use AI to process tables.
            process_images_with_ai (bool, optional): Use AI to process images.

        Returns:
            None
        """

        source_file: Path = Path(input_path)
        is_docx: bool = source_file.suffix.lower() == ".docx"
        is_txt: bool = source_file.suffix.lower() == ".txt"

        if is_txt:
            shutil.copy2(str(source_file), str(output_path))
            return

        if steps not in {1, 2, 3, 4, 5, 6}:
            raise ValueError(f"Invalid step value: {steps}. Must be between 1 and 6.")

        if steps == 6:
            steps_range = range(1, 6)  # 6: All steps in sequence
        else:
            steps_range = range(1, steps + 1)  # E.g. 3 -> 1,2,3; 5 -> 1,2,3,4,5

        for step in steps_range:
            input_stem = source_file.stem
            if step == 1:
                # 1: Document cleaning (header, footer, table of contents, empty lines/pages, abbreviations, footnotes)
                if is_docx:
                    images_subdir = Path(self.images_dir) / input_stem
                    FileService.extract_images_from_docx(
                        source_file=source_file,
                        images_dir=images_subdir,
                    )

                    cleaner = DOCXDocumentCleaner(source_file=source_file)
                    if remove_headers:
                        cleaner.remove_headers()
                    if remove_footers:
                        cleaner.remove_footers()
                    if remove_toc:
                        cleaner.remove_toc()
                    if remove_empty:
                        cleaner.remove_empty()
                    if abbreviation_strategy != "none":
                        abbreviations = cleaner.abbreviation_extractors(
                            abbreviations_output=Path(self.temp_dir)
                            / f"{input_stem}_abbreviations.json"
                        )
                        if abbreviations:
                            cleaner.abbreviation_modifiers(
                                abbreviations=abbreviations,
                                strategy=abbreviation_strategy,
                            )
                    if footnote_handling != "none":
                        cleaner.footnote_handling(strategy=footnote_handling)
                else:
                    pass  # No document cleaning step for non-.docx files

            elif (
                step == 2
            ):  # 2: Convert document (PDF/DOCX) to Markdown format (Docling)
                images_subdir = Path(self.images_dir) / input_stem
                docling_converter = DoclingConverter(
                    source_file=source_file,
                    markdown_file=Path(self.temp_dir) / f"{input_stem}_markdown.md",
                    images_dir=images_subdir,
                )
                supported_extensions = [".docx", ".pdf"]
                docling_converter.validate_document(
                    supported_extensions=supported_extensions
                )
                docling_converter.docling_process_document()
                docling_converter.clean_markdown_file()
                docling_converter.replace_image_placeholders_with_markdown()

            elif step == 3:
                # 3: Process Markdown file to JSON format (Unstructured)
                unstructured_converter = UnstructuredConverter(
                    source_file=source_file,
                    markdown_file=Path(self.temp_dir) / f"{input_stem}_markdown.md",
                    json_file=Path(self.temp_dir) / f"{input_stem}_json.json",
                )
                unstructured_converter.unstructured_process_markdown()

            elif step == 4:
                # 4: Enrich JSON content with table and image summaries (OpenAI)
                enrichment = Enrichment(
                    source_file=source_file,
                    json_file=Path(self.temp_dir) / f"{input_stem}_json.json",
                    client=client,
                    enriched_json_file=Path(self.temp_dir)
                    / f"{input_stem}_enriched_json.json",
                )
                enrichment.enrich_json(
                    process_table_with_ai=process_table_with_ai,
                    process_images_with_ai=process_images_with_ai,
                    text_model=self.text_model,
                    image_model=self.image_model,
                )

            elif step == 5:
                export = Export(
                    enriched_json_file=Path(self.temp_dir)
                    / f"{input_stem}_enriched_json.json",
                    txt_file=Path(self.output_dir)
                    / f"{input_stem.replace('working_', '')}.txt",
                )
                export.export_text_from_enriched_json()


class DOCXDocumentCleaner:
    """
    Cleans and preprocesses DOCX documents for downstream processing.

    This class provides comprehensive cleaning and normalization for DOCX files, including header/footer removal, table of contents (TOC) removal, empty paragraph deletion, abbreviation extraction and modification, and footnote handling. Designed for flexible integration into document processing pipelines.

    Args:
        source_file (Path): Path to the source DOCX file.
    """

    def __init__(self, source_file: Path) -> None:
        self.source_file = source_file
        self.doc = Document(str(self.source_file))

    def remove_headers(self) -> None:
        """
        Remove all headers from the document.

        Iterates through document sections and clears all paragraphs in headers.

        Returns:
            None
        """
        for section in self.doc.sections:
            for header in section.header.paragraphs:
                header.clear()
        self.doc.save(str(self.source_file))

    def remove_footers(self) -> None:
        """
        Remove all footers from the document.

        Iterates through document sections and clears all paragraphs in footers.

        Returns:
            None
        """
        for section in self.doc.sections:
            for footer in section.footer.paragraphs:
                footer.clear()
        self.doc.save(str(self.source_file))

    def remove_toc(self) -> None:
        """
        Remove the table of contents (TOC) from the document using multiple strategies.

        Applies XML-based, field code-based, paragraph-based, text-based, table-based, and original logic methods.

        Returns:
            None
        """

        def remove_toc_by_xml(doc: Document) -> bool:
            """
            Removes the table of contents (TOC) from a Word document based on its XML structure.

            Args:
                doc (Document): The python-docx Document object.

            Returns:
                bool: True if any TOC was removed, otherwise False.
            """
            document = doc.element.body
            paragraphs_to_delete = []
            toc_pattern = re.compile(r"TOC.*")
            pageref_pattern = re.compile(r"PAGEREF _Toc\d+.*")
            toc_hyperlink_pattern = re.compile(r"_TOC_\d+")

            toc_found = False
            paragraph_index = 0

            for paragraph in document.xpath(".//w:p"):
                is_toc = False
                paragraph_index += 1

                instr_texts = paragraph.xpath(".//w:instrText")
                for instr in instr_texts:
                    if instr.text and (
                        toc_pattern.match(instr.text)
                        or pageref_pattern.match(instr.text)
                    ):
                        is_toc = True
                        paragraphs_to_delete.append(paragraph)
                        break

                hyperlinks = paragraph.xpath(".//w:hyperlink")
                for hyperlink in hyperlinks:
                    anchor = hyperlink.get(qn("w:anchor"))
                    if anchor and toc_hyperlink_pattern.match(anchor):
                        is_toc = True
                        paragraphs_to_delete.append(paragraph)
                        break

                style = paragraph.xpath(".//w:pStyle/@w:val")
                if style and style[0] == "Listaszerbekezds" and toc_found:
                    is_toc = True
                    paragraphs_to_delete.append(paragraph)
                elif style and any(s.startswith(("TJ", "TOC", "toc")) for s in style):
                    is_toc = True
                    paragraphs_to_delete.append(paragraph)

                if is_toc and not toc_found:
                    toc_found = True

                elif toc_found:
                    if style and (
                        style[0].startswith("Heading")
                        or (
                            not style[0]
                            .lower()
                            .startswith(("tj", "toc", "tartalomjegyzk"))
                            and style[0] != "Listaszerbekezds"
                        )
                    ):
                        text_elements = paragraph.xpath(".//w:t")
                        if text_elements and any(t.text.strip() for t in text_elements):
                            toc_found = False
                            break

            sdt_elements = document.xpath(".//w:sdt")
            for sdt in sdt_elements:
                sdt_pr = sdt.find(qn("w:sdtPr"))
                if sdt_pr is not None:
                    doc_part_obj = sdt_pr.find(qn("w:docPartObj"))
                    if doc_part_obj is not None:
                        tag_elem = sdt_pr.find(qn("w:tag"))
                        if (
                            tag_elem is not None
                            and "toc" in tag_elem.get(qn("w:val"), "").lower()
                        ):
                            parent = sdt.getparent()
                            if parent is not None:
                                parent.remove(sdt)
                                continue

                        doc_part = doc_part_obj.find(qn("w:docPartGallery"))
                        if (
                            doc_part is not None
                            and doc_part.get(qn("w:val")) == "Table of Contents"
                        ):
                            parent = sdt.getparent()
                            if parent is not None:
                                parent.remove(sdt)
                                continue

                            for para in sdt.xpath(".//w:p"):
                                paragraphs_to_delete.append(para)

            for paragraph in paragraphs_to_delete:
                try:
                    parent = paragraph.getparent()
                    if parent is not None:
                        parent.remove(paragraph)
                except Exception as e:
                    print(f"Hiba a bekezdés törlése közben: {e}")

            for paragraph in document.xpath(".//w:p"):
                if not paragraph.xpath(".//w:t") and not paragraph.xpath(
                    ".//w:instrText"
                ):
                    try:
                        parent = paragraph.getparent()
                        if parent is not None:
                            parent.remove(paragraph)
                    except Exception as e:
                        print(f"Hiba az üres bekezdés törlésekor: {e}")

            return len(paragraphs_to_delete) > 0

        def remove_toc_by_field(doc: Document) -> bool:
            """
            Removes the table of contents (TOC) from a Word document based on field codes.

            Args:
                doc (Document): The python-docx Document object.

            Returns:
                bool: True if any TOC was removed, otherwise False.
            """
            toc_removed = False
            paragraphs_to_delete = []

            for para_index, para in enumerate(doc.paragraphs):
                for run in para.runs:
                    try:
                        for elem in run._element:
                            if (
                                (
                                    elem.tag == qn("w:fldChar")
                                    and "TOC" in elem.getparent().xml
                                )
                                or (
                                    elem.tag == qn("w:instrText")
                                    and elem.text
                                    and ("TOC" in elem.text or "PAGEREF" in elem.text)
                                )
                                or (elem.getparent().tag == qn("w:sdt"))
                            ):
                                paragraphs_to_delete.append(para)
                                toc_removed = True
                                break
                    except Exception as elem_error:
                        print(
                            f"Hiba az elemek keresése során a {para_index}. bekezdésben: {str(elem_error)}"
                        )

                try:
                    sdt_tag = qn("w:sdt")
                    sdt = para._element.find(sdt_tag)
                    if sdt is not None:
                        nsmap = {
                            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        }
                        sdt_pr_tag = qn("w:sdtPr")
                        sdt_pr = sdt.find(sdt_pr_tag, namespaces=nsmap)
                        if sdt_pr is not None:
                            doc_part_obj = sdt_pr.find(
                                qn("w:docPartObj"), namespaces=nsmap
                            )
                            if doc_part_obj is not None:
                                doc_part = doc_part_obj.find(
                                    qn("w:docPartGallery"), namespaces=nsmap
                                )
                                if (
                                    doc_part is not None
                                    and doc_part.get("val") == "Table of Contents"
                                ):
                                    paragraphs_to_delete.append(para)
                                    toc_removed = True
                except Exception as sdt_error:
                    print(
                        f"Hiba az SDT keresése során a {para_index}. bekezdésben: {str(sdt_error)}"
                    )

            for para_index, para in enumerate(paragraphs_to_delete):
                try:
                    if para._element.getparent() is not None:
                        para._element.getparent().remove(para._element)
                except Exception as e:
                    print(f"Hiba a bekezdés törlésekor: {str(e)}")

            return toc_removed

        def remove_toc_by_paragraphs(doc: Document) -> bool:
            """
            Removes the table of contents (TOC) from a Word document based on paragraph styles and content.

            Args:
                doc (Document): The python-docx Document object.

            Returns:
                bool: True if any TOC was removed, otherwise False.
            """
            toc_found = False
            paragraphs_to_delete = []

            for para in doc.paragraphs:
                is_toc = (
                    para.style.name.lower().startswith(("toc", "tartalomjegyzk", "tj"))
                    or para.style.name
                    in ["TJ1", "TJ2", "TJ3", "TJ4", "TJ5", "Jegyzkhivatkozs"]
                    or "tartalomjegyzék" in para.text.lower()
                    or "Table of Contents" in para.text.lower()
                    or any("TOC" in run.text.lower() for run in para.runs)
                    or any("PAGEREF" in run.text for run in para.runs)
                    or any(
                        run._element.find(qn("w:hyperlink")) is not None
                        for run in para.runs
                    )
                    or any(
                        elem.tag == qn("w:fldChar")
                        for run in para.runs
                        for elem in run._element
                    )
                    or any(
                        "......" in run.text or "." * 10 in run.text
                        for run in para.runs
                    )
                    or (
                        any(run.text.isdigit() for run in para.runs)
                        and any(
                            run._element.find(qn("w:tab")) is not None
                            for run in para.runs
                        )
                    )
                )

                if (
                    para.style.name == "Listaszerbekezds"
                    and para._element.find(qn("w:numPr")) is not None
                    and toc_found
                ):
                    is_toc = True

                if is_toc:
                    toc_found = True
                    paragraphs_to_delete.append(para)
                elif toc_found:
                    if (
                        para.style.name.startswith("Heading")
                        and para.text.strip()
                        or not para.style.name.lower().startswith(
                            ("tj", "toc", "tartalomjegyzk")
                        )
                        and para.text.strip()
                        and para.style.name != "Listaszerbekezds"
                    ):
                        break

            for para in paragraphs_to_delete:
                try:
                    p = para._element
                    if p.getparent() is not None:
                        p.getparent().remove(p)
                except Exception as e:
                    print(f"Hiba a bekezdés törlésekor: {str(e)}")

            return len(paragraphs_to_delete) > 0

        def remove_toc_by_text(doc: Document) -> bool:
            """
            Removes the table of contents (TOC) from a Word document using keyword-based text search.

            Args:
                doc (Document): The python-docx Document object.

            Returns:
                bool: True if any TOC was removed, otherwise False.
            """
            paragraphs_to_delete = []
            toc_found = False

            for para in doc.paragraphs:
                if (
                    "tartalomjegyzék" in para.text.lower()
                    or "Table of Contents" in para.text.lower()
                ):
                    toc_found = True
                    paragraphs_to_delete.append(para)
                elif toc_found:
                    if (
                        para.style.name
                        in ["TJ1", "TJ2", "TJ3", "TJ4", "TJ5", "Jegyzkhivatkozs", "TOC"]
                        or para.style.name.lower().startswith(
                            ("toc", "tartalomjegyzk", "TJ")
                        )
                        or any("PAGEREF" in run.text for run in para.runs)
                        or any(
                            run._element.find(qn("w:hyperlink")) is not None
                            for run in para.runs
                        )
                        or any(
                            "......" in run.text or "." * 10 in run.text
                            for run in para.runs
                        )
                    ):
                        paragraphs_to_delete.append(para)
                    elif (
                        para.style.name.startswith("Heading")
                        and para.text.strip()
                        or para.style.name
                        not in [
                            "TJ1",
                            "TJ2",
                            "TJ3",
                            "TJ4",
                            "TJ5",
                            "Jegyzkhivatkozs",
                            "TOC",
                        ]
                        or para.style.name.lower().startswith(
                            ("toc", "tartalomjegyzk", "TJ")
                        )
                        and para.text.strip()
                    ):
                        break

            for para in paragraphs_to_delete:
                p = para._element
                p.getparent().remove(p)

            return len(paragraphs_to_delete) > 0

        def remove_toc_by_table(doc: Document) -> bool:
            """
            Removes the table of contents (TOC) from a Word document based on tables.

            Args:
                doc (Document): The python-docx Document object.

            Returns:
                bool: True if any TOC was removed, otherwise False.
            """
            toc_removed = False
            for table in doc.tables:
                if any(
                    "......" in cell.text or "." * 10 in cell.text
                    for row in table.rows
                    for cell in row.cells
                ):
                    table._element.getparent().remove(table._element)
                    toc_removed = True
            return toc_removed

        def remove_toc_original(doc: Document) -> bool:
            """
            Removes the table of contents (TOC) from a Word document based on the "tartalomjegyzék" keyword.

            Args:
                doc (Document): The python-docx Document object from which the TOC should be removed.

            Returns:
                bool: True if any TOC was removed, otherwise False.
            """
            toc_heading = None
            toc_start = None
            toc_end = None
            toc_removed = False

            for i, para in enumerate(doc.paragraphs):
                if (
                    "tartalomjegyzék" in para.text.lower()
                    or "table of contents" in para.text.lower()
                ):
                    toc_heading = para
                    toc_start = i
                    break

            if toc_start is not None:
                for i in range(toc_start + 1, len(doc.paragraphs)):
                    para = doc.paragraphs[i]
                    if not para.text.strip() or any(
                        run.text == "\x0c" for run in para.runs
                    ):
                        toc_end = i
                        break
                    elif para.style.name.startswith("Heading"):
                        toc_end = i
                        break

                if toc_end is not None:
                    for i in range(toc_start, toc_end + 1):
                        if toc_start < len(doc.paragraphs):
                            doc.paragraphs[toc_start]._element.getparent().remove(
                                doc.paragraphs[toc_start]._element
                            )
                            toc_removed = True
                else:
                    if toc_heading is not None:
                        toc_heading._element.getparent().remove(toc_heading._element)
                        toc_removed = True

            return toc_removed

        methods = [
            (remove_toc_by_xml, "XML-based"),
            (remove_toc_by_field, "Field code-based"),
            (remove_toc_by_paragraphs, "Paragraph-based"),
            (remove_toc_by_text, "Text search-based"),
            (remove_toc_by_table, "Table-based"),
            (remove_toc_original, "Original removal logic"),
        ]
        for method, method_name in methods:
            try:
                if method(self.doc):
                    continue
            except Exception as e:
                raise RuntimeError(f"Error during {method_name} method: {e}")

        self.doc.save(str(self.source_file))

    def remove_empty(self) -> None:
        """
        Remove empty paragraphs from the document, preserving page layout and orientation.

        Checks if a paragraph is empty (text, runs, objects, sections) and removes it.

        Returns:
            None
        """

        def is_paragraph_empty(paragraph: Paragraph) -> bool:
            """
            Checks if a paragraph is empty.

            Args:
                paragraph (Paragraph): The paragraph to check.

            Returns:
                bool: True if the paragraph is empty, otherwise False.
            """

            if paragraph.text.strip():
                return False

            for run in paragraph.runs:
                if run.text.strip():
                    return False

            if len(paragraph._element.xpath(".//w:drawing|.//w:pict|.//w:object")) > 0:
                return False

            if paragraph._element.xpath(".//w:sectPr"):
                return False

            return True

        empty_paragraphs = [
            para for para in self.doc.paragraphs if is_paragraph_empty(para)
        ]

        for para in empty_paragraphs:
            parent = para._element.getparent()

            if parent.tag.endswith("sectPr"):
                continue

            try:
                parent.remove(para._element)
            except Exception as e:
                print(f"Could not remove paragraph: {str(e)}")
                continue

        self.doc.save(str(self.source_file))

    def abbreviation_extractors(
        self, abbreviations_output: str
    ) -> Dict[str, List[str]]:
        """
        Extract abbreviations and their definitions from the document.

        Stores multiple definitions per abbreviation and saves the result to a JSON file.

        Args:
            abbreviations_output (str): Path to save the extracted abbreviations.

        Returns:
            Dict[str, List[str]]: Dictionary of abbreviations and their definitions.
        """

        def is_uppercase_abbr(abbr: str) -> bool:
            """
            Checks if the abbreviation is uppercase or a short uppercase form (e.g., MNB, ABCDE).

            Args:
                abbr (str): The abbreviation to check.

            Returns:
                bool: True if the abbreviation is uppercase or a short uppercase form, otherwise False.
            """

            if not abbr:
                return False
            return abbr.isupper() or (
                abbr[0].isupper() and len(abbr) <= 5 and abbr.isalpha()
            )

        def is_law_or_decree(text: str) -> bool:
            """
            Checks if the text refers to a law or decree.

            Args:
                text (str): The text to check.

            Returns:
                bool: True if the text refers to a law or decree, otherwise False.
            """
            patterns = [
                r"\d{1,4}/\d{4}\.\s*\([IVXLCDM]+\.\s*\d+\.\)",  # e.g., 87/2015. (IV. 9.)
                r"\d{4}\.\s*évi\s+[IVXLCDM]+\.",  # e.g., 2011. évi CCIV.
                r"Korm\.\s*rendelet",  # Government decree
                r"törvény",  # Law
            ]
            return any(
                re.search(pattern, text, flags=re.IGNORECASE) for pattern in patterns
            )

        def get_hungarian_initial(word: str) -> str:
            """
            Extracts the initial letter from a Hungarian word, considering digraphs and trigraphs.

            Args:
                word (str): The word to analyze.

            Returns:
                str: The initial letter, digraph, or trigraph.
            """
            if not word:
                return ""

            word = word.lower()
            hungarian_trigraphs = ["dzs"]
            hungarian_digraphs = ["cs", "dz", "gy", "ly", "ny", "sz", "ty", "zs"]

            for trigraph in hungarian_trigraphs:
                if word.startswith(trigraph):
                    return trigraph.upper()

            for digraph in hungarian_digraphs:
                if word.startswith(digraph):
                    return digraph.upper()

            return word[0].upper() if word[0].isalpha() else ""

        def validate_uppercase_abbr(abbr: str, words: List[str]) -> Tuple[bool, str]:
            """
            Validates uppercase abbreviations based on initials.

            Args:
                abbr (str): The abbreviation.
                words (List[str]): The words of the definition.

            Returns:
                Tuple[bool, str]: (validity, final definition).
            """
            if not abbr:
                return False, ""

            abbr = abbr.upper()
            full_definition = " ".join(words)

            filtered_words = [w for w in words if w.lower() not in ["és", "a", "az"]]
            filtered_indices = [
                i for i, w in enumerate(words) if w.lower() not in ["és", "a", "az"]
            ]

            if is_law_or_decree(full_definition):
                return True, full_definition

            best_match_len = 0
            best_definition = None
            best_sub_words_len = float("inf")

            for i in range(len(filtered_words)):
                sub_words = filtered_words[i:]
                sub_initials = [
                    get_hungarian_initial(w)
                    for w in sub_words
                    if get_hungarian_initial(w)
                ]
                sub_initial_str = "".join(sub_initials)

                if len(sub_initial_str) >= len(abbr) and sub_initial_str.startswith(
                    abbr
                ):
                    start_index = filtered_indices[i]
                    definition = " ".join(words[start_index:])
                    return True, definition

                for match_len in range(len(abbr), 0, -1):
                    if len(sub_initials) >= match_len:
                        for start in range(len(sub_initials) - match_len + 1):
                            if all(
                                sub_initials[start + j] == abbr[j]
                                for j in range(match_len)
                            ):
                                start_index = filtered_indices[i]
                                max_words = min(5, len(words[start_index:]))
                                definition = " ".join(
                                    words[start_index : start_index + max_words]
                                )
                                if match_len > best_match_len or (
                                    match_len == best_match_len
                                    and len(sub_words) < best_sub_words_len
                                ):
                                    best_match_len = match_len
                                    best_definition = definition
                                    best_sub_words_len = len(sub_words)

            if best_definition and best_match_len >= 1:
                return True, best_definition

            abbr_len = get_adjusted_abbr_len(abbr)
            word_count = len(filtered_words)
            max_diff = min(6, len(abbr) * 2)

            if abs(abbr_len - word_count) <= max_diff:
                max_words = min(5, len(words))
                definition = " ".join(words[-max_words:])
                return True, definition

            return False, " ".join(words)

        def validate_lowercase_abbr(abbr: str, words: List[str]) -> Tuple[bool, str]:
            """
            Validates lowercase abbreviations based on the definition.

            Args:
                abbr (str): The abbreviation.
                words (List[str]): The words of the definition.

            Returns:
                Tuple[bool, str]: (validity, final definition).
            """
            if not abbr:
                return False, ""

            full_definition = " ".join(words)

            normalized_abbr = abbr.lower()
            filtered_words = [w for w in words if w.lower() not in ["és", "a", "az"]]
            filtered_indices = [
                i for i, w in enumerate(words) if w.lower() not in ["és", "a", "az"]
            ]

            if is_law_or_decree(full_definition):
                return True, full_definition

            abbr_words = normalized_abbr.split()
            max_words = (
                len(abbr_words) * 4 if len(abbr_words) == 1 else len(abbr_words) * 3
            )
            best_definition = None
            best_match_len = 0

            for i in range(len(filtered_words) - len(abbr_words) + 1):
                sub_words = filtered_words[i : i + len(abbr_words)]

                if all(
                    sub_words[j].lower() == abbr_words[j]
                    for j in range(len(abbr_words))
                ):
                    start_index = max(
                        0, filtered_indices[i] - (max_words - len(abbr_words))
                    )
                    end_index = min(len(words), start_index + max_words)
                    definition_words = words[start_index:end_index]
                    definition = " ".join(definition_words)
                    match_len = len(abbr_words)

                    if match_len > best_match_len:
                        best_match_len = match_len
                        best_definition = definition

            if best_definition:
                return True, best_definition

            definition = " ".join(words[max(0, len(words) - max_words) :])
            return True, definition

        def normalize_text(text: str) -> str:
            """
            Normalizes the text by removing Hungarian prefixes and suffixes.

            Args:
                text (str): The input text.

            Returns:
                str: The normalized text.
            """
            prefixes = ["az ", "a ", "Az ", "A "]
            suffixes = [
                "ának",
                "anak",
                "enek",
                "nak",
                "nek",
                "ban",
                "ben",
                "bol",
                "ből",
                "re",
                "val",
                "vel",
                "ert",
                "ig",
                "hoz",
                "hez",
                "höz",
                "ul",
                "ül",
                "ok",
                "ek",
            ]

            for prefix in prefixes:
                if text.lower().startswith(prefix):
                    text = text[len(prefix) :].strip()
                    break

            for suffix in suffixes:
                if text.lower().endswith(suffix):
                    text = text[: -len(suffix)].strip()
                    break

            return text

        def get_adjusted_abbr_len(abbr: str) -> int:
            """
            Calculates the number of units in the abbreviation, considering digraphs and trigraphs.

            Args:
                abbr (str): The abbreviation.

            Returns:
                int: The number of units.
            """
            abbr = abbr.lower()
            hungarian_trigraphs = ["dzs"]
            hungarian_digraphs = ["cs", "dz", "gy", "ly", "ny", "sz", "ty", "zs"]

            units = []
            i = 0
            while i < len(abbr):
                if i + 3 <= len(abbr) and abbr[i : i + 3] in hungarian_trigraphs:
                    units.append(abbr[i : i + 3])
                    i += 3
                elif i + 2 <= len(abbr) and abbr[i : i + 2] in hungarian_digraphs:
                    units.append(abbr[i : i + 2])
                    i += 2
                else:
                    units.append(abbr[i])
                    i += 1
            return len(units)

        def validate_abbreviation(abbr: str, full_text: str) -> Tuple[bool, str]:
            """
            Validates whether the abbreviation matches the definition.

            Args:
                abbr (str): The abbreviation.
                full_text (str): The full definition text.

            Returns:
                Tuple[bool, str]: (validity, final definition).
            """
            definition = normalize_text(full_text)
            words = definition.split()

            if is_uppercase_abbr(abbr):
                return validate_uppercase_abbr(abbr, words)
            else:
                return validate_lowercase_abbr(abbr, words)

        def split_abbreviations(abbrs: str) -> List[str]:
            """
            Splits abbreviations by comma and cleans them.

            Args:
                abbrs (str): The abbreviations string (comma-separated).

            Returns:
                List[str]: List of cleaned abbreviations.
            """
            if not abbrs:
                return []
            return [
                abbr.strip().rstrip(".")
                for abbr in re.split(r",\s*|\svagy\s|/\s*", abbrs)
                if abbr.strip()
            ]

        def process_abbreviations(
            full_text: str, abbrs: str, abbreviation_dict: Dict[str, List[str]]
        ) -> None:
            """
            Processes abbreviations, including handling multiple abbreviations.

            Args:
                full_text (str): The full definition text.
                abbrs (str): The abbreviations (comma-separated).
                abbreviation_dict (Dict[str, List[str]]): The abbreviation-definition dictionary.
            """
            abbr_list = split_abbreviations(abbrs)

            for abbr in abbr_list:
                if not abbr:
                    continue
                is_valid, definition = validate_abbreviation(abbr, full_text)
                if is_valid:
                    if abbr not in abbreviation_dict:
                        abbreviation_dict[abbr] = []
                    if definition not in abbreviation_dict[abbr]:
                        abbreviation_dict[abbr].append(definition)

        def extract_abbreviations(doc: Document) -> Dict[str, List[str]]:
            """
            Extracts abbreviations and their definitions from the document.

            Args:
                doc (Document): The python-docx Document object.

            Returns:
                Dict[str, List[str]]: Dictionary of abbreviations and their definitions.
            """
            abbreviation_dict: Dict[str, List[str]] = {}
            pattern = re.compile(
                r"(.+?)\s*\((?:(?:a\s*)?továbbiakban(?:\s*együtt)?\s*:\s*|továbbiakban\s+)(.+?)\)",
                flags=re.IGNORECASE,
            )

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                matches = pattern.findall(para.text)
                for full_text, abbrs in matches:
                    process_abbreviations(
                        full_text.strip(), abbrs.strip(), abbreviation_dict
                    )

            return abbreviation_dict

        def save_abbreviations(abbreviations: Dict[str, List[str]], path: str) -> None:
            """
            Saves the abbreviation dictionary to a JSON file.

            Args:
                abbreviations (Dict[str, List[str]]): The abbreviation dictionary.
                path (str): The file path to save the dictionary.

            Raises:
                RuntimeError: If saving the abbreviation dictionary fails.
            """

            try:
                FileService.write_json(abbreviations, path)
            except Exception as e:
                raise RuntimeError(f"Error saving abbreviation dictionary: {str(e)}")

        abbreviations = extract_abbreviations(self.doc)
        save_abbreviations(abbreviations, abbreviations_output)

        return abbreviations

    def abbreviation_modifiers(
        self, abbreviations: Dict[str, List[str]], strategy: str
    ) -> None:
        """
        Modify abbreviations in the document according to the specified strategy.

        Args:
            abbreviations (Dict[str, List[str]]): Dictionary of abbreviations and definitions.
            strategy (str): Modification strategy, either "inline" or "section".

        Returns:
            None
        """

        def remove_abbreviation_phrases(
            doc: Document, abbreviations: Dict[str, List[str]]
        ) -> None:
            """
            Removes the phrase '(a továbbiakban: XYZ)' or '(továbbiakban XYZ)' from the document,
            keeping the abbreviation in parentheses if XYZ is found in the abbreviation dictionary.

            Args:
                doc (Document): The python-docx Document object.
                abbreviations (Dict[str, List[str]]): The abbreviation dictionary.

            Returns:
                None
            """

            pattern = re.compile(
                r"(.+?)\s*\((?:(?:a\s*)?továbbiakban(?:\s*együtt)?\s*:\s*|továbbiakban\s+)(.+?)\)",
                flags=re.IGNORECASE,
            )

            for para in doc.paragraphs:
                new_text = para.text
                matches = pattern.finditer(new_text)

                for match in matches:
                    full_match = match.group(0)
                    full_expression = match.group(1).strip()
                    abbreviation = match.group(2).strip()

                    if abbreviation in abbreviations or any(
                        abbr in abbreviations for abbr in abbreviation.split(", ")
                    ):
                        new_text = re.sub(
                            re.escape(full_match),
                            f"{full_expression} ({abbreviation})",
                            new_text,
                            flags=re.IGNORECASE,
                        )

                para.text = new_text.strip()

        def insert_abbreviations(
            doc: Document, abbr_dict: Dict[str, List[str]]
        ) -> None:
            """
            Inserts abbreviation definitions where the abbreviation appears alone,
            without a parenthetical definition and not as part of an existing parenthetical structure.

            Args:
                doc (Document): The python-docx Document object.
                abbr_dict (Dict[str, List[str]]): The abbreviation dictionary.

            Returns:
                None
            """
            inserted_count = 0

            for para in doc.paragraphs:
                para_text = para.text
                modified = False
                new_text = para_text
                replacements = []

                for abbr, definitions in abbr_dict.items():
                    if not definitions:
                        continue
                    full = definitions[0]
                    pattern = re.compile(
                        rf"(?<!\w)\b{re.escape(abbr)}\b(?!\w)(?!\s*\([^)]*\))",
                        flags=re.UNICODE | re.IGNORECASE,
                    )

                    for match in pattern.finditer(para_text):
                        start, end = match.span()
                        abbr_text = match.group(0)

                        before_text = para_text[:start]
                        open_paren_count = before_text.count("(") - before_text.count(
                            ")"
                        )
                        if open_paren_count > 0:
                            continue

                        if (start > 0 and para_text[start - 1].isalnum()) or (
                            end < len(para_text) and para_text[end].isalnum()
                        ):
                            continue

                        replacements.append((start, end, f"{abbr_text} ({full})"))

                if replacements:
                    for start, end, replacement in sorted(replacements, reverse=True):
                        new_text = new_text[:start] + replacement + new_text[end:]
                        modified = True
                        inserted_count += 1

                if modified:
                    para.text = new_text

        def prepend_abbreviation_section(
            doc: Document, abbr_dict: Dict[str, List[str]]
        ) -> None:
            """
            Lists abbreviations and their definitions at the beginning of the document as new paragraphs, in alphabetical order, including all definitions.

            Args:
                doc (Document): The python-docx Document object.
                abbr_dict (Dict[str, List[str]]): The abbreviation dictionary.

            Returns:
                None
            """

            if not abbr_dict:
                return

            section_title = doc.paragraphs[0].insert_paragraph_before("Rövidítések")
            try:
                section_title.style = "Heading 2"
            except KeyError:
                style = doc.styles.add_style("Heading 2", 1)
                font = style.font
                font.name = "Times New Roman"
                font.size = Pt(14)
                font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
                style.paragraph_format.space_before = Pt(18)
                style.paragraph_format.space_after = Pt(0)
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                section_title.style = style

            try:
                normal_style = doc.styles["Normal"]
            except KeyError:
                normal_style = doc.styles.add_style("Normal", 1)
                normal_style.font.name = "Times New Roman"
                normal_style.font.size = Pt(12)
                normal_style.paragraph_format.space_before = Pt(6)
                normal_style.paragraph_format.space_after = Pt(6)

            sorted_abbrs = sorted(abbr_dict.items(), key=lambda x: x[0].lower())

            for abbr, definitions in sorted_abbrs:
                for i, full in enumerate(definitions, 1):
                    para_text = (
                        f"{abbr}: {full}" if i == 1 else f"{abbr} ({i}.): {full}"
                    )
                    para = doc.paragraphs[0].insert_paragraph_before(para_text)
                    para.style = normal_style

            doc.paragraphs[0].insert_paragraph_before("")

        remove_abbreviation_phrases(self.doc, abbreviations)
        if strategy == "inline":
            insert_abbreviations(self.doc, abbreviations)
        elif strategy == "section":
            prepend_abbreviation_section(self.doc, abbreviations)

        self.doc.save(str(self.source_file))

    def footnote_handling(self, strategy: str) -> None:
        """
        Handle footnotes in the document according to the specified strategy.

        Args:
            strategy (str): Footnote handling strategy, either "remove" or "insert".

        Returns:
            None
        """

        def remove_footnote_references(doc: Document) -> None:
            """
            Removes footnote references and distracting elements based on regular patterns.

            Args:
                doc (Document): The python-docx Document object.

            Returns:
                None
            """

            for para in doc.paragraphs:
                inline_elements = para._element.xpath(".//w:footnoteReference")
                for el in inline_elements:
                    el.getparent().remove(el)

            patterns = [
                r"\[\*\]",  # [*]
                r"\^[0-9]+",  # ^1, ^2
            ]

            for para in doc.paragraphs:
                original = para.text
                cleaned = original
                for pattern in patterns:
                    cleaned = re.sub(pattern, "", cleaned)
                para.text = cleaned

        def insert_footnotes(doc: Document) -> None:
            """
            Inserts the content of footnotes into the document.

            Args:
                doc (Document): The python-docx Document object.

            Returns:
                None
            """

            footnote_part = None

            for rel in doc.part.rels.values():
                if rel.reltype == RT.FOOTNOTES:
                    footnote_part = rel._target
                    break

            if not footnote_part:
                return

            footnotes_tree = ET.fromstring(footnote_part.blob)
            nsmap = footnotes_tree.tag.split("}")[0].strip("{")

            footnote_map = {
                fn.attrib.get(f"{{{nsmap}}}id"): "".join(
                    node.text or "" for node in fn.iter() if node.text
                )
                for fn in footnotes_tree.findall(
                    ".//w:footnote", namespaces={"w": nsmap}
                )
            }

            for para in doc.paragraphs:
                for run in para.runs:
                    r_el = run._element
                    footnote_refs = r_el.xpath('.//*[local-name()="footnoteReference"]')
                    for ref in footnote_refs:
                        fid = ref.attrib.get(f"{{{nsmap}}}id")
                        footnote_text = footnote_map.get(fid)
                        if footnote_text:
                            placeholder = f"<<lábjegyzet: {footnote_text.strip()}>>"
                            run.text += f" {placeholder}"
                        parent = ref.getparent()
                        if parent is not None:
                            parent.remove(ref)

        if strategy == "remove":
            remove_footnote_references(self.doc)
        elif strategy == "insert":
            insert_footnotes(self.doc)

        self.doc.save(str(self.source_file))


class DoclingConverter:
    """
    Converts documents to Markdown format and manages image references for downstream processing.

    This class uses the Docling library to convert DOCX/PDF documents to Markdown, applies normalization and cleaning steps, and replaces image placeholders with correct Markdown syntax. The images directory is configurable for flexible integration with different pipeline setups.

    Args:
        source_file (Path): Path to the source document file.
        markdown_file (Path): Path to the output Markdown file.
        images_dir (str): Path to the images directory (where extracted images are stored).
    """

    def __init__(self, source_file: Path, markdown_file: Path, images_dir: str) -> None:
        self.source_file = source_file
        self.markdown_file = markdown_file
        self.images_dir = images_dir

    def validate_document(self, supported_extensions: list[str]) -> bool:
        """
        Checks if the document file exists and is in a supported format.

        The function performs two basic checks:
        1. Checks if the file physically exists in the filesystem.
        2. Checks if the file extension is in the supported_extensions list (case-insensitive).

        Args:
            supported_extensions (list[str]): List of supported file extensions (e.g., ['.docx', '.pdf']).

        Returns:
            bool: True if the file exists and is in a supported format, otherwise False.
        """

        if not self.source_file.exists():
            return False

        if self.source_file.suffix.lower() not in supported_extensions:
            return False

        return True

    def docling_process_document(self) -> None:
        """
        Processes a document into Markdown format using Docling.

        Uses the Docling library to convert the document, then saves the generated
        Markdown content to the specified file.

        Raises:
            RuntimeError: If document conversion fails.
        """

        converter = DocumentConverter()

        try:
            result = converter.convert(self.source_file)
        except Exception as e:
            raise RuntimeError(f"Document conversion failed: {str(e)}")

        if not result.document:
            raise RuntimeError("Document conversion failed")

        markdown_content = result.document.export_to_markdown()

        FileService.write_text(markdown_content, self.markdown_file)

    def clean_markdown_file(self) -> None:
        """
        Cleans and normalizes a Markdown file in place, applying all fixers.

        Returns:
            None
        """

        def fix_markdown_tables(lines: list[str]) -> list[str]:
            """
            Ensures that Markdown tables have blank lines around them and correct syntax.

            Args:
                lines (list[str]): Lines of the Markdown file.

            Returns:
                list[str]: Modified lines.
            """
            new_lines = []
            i = 0
            while i < len(lines):
                line = lines[i]
                if "|" in line and re.match(r"^\s*\|?(.+\|)+\s*$", line):
                    if i + 1 < len(lines) and re.match(
                        r"^\s*\|?\s*:?[-| ]+:?\s*\|?\s*$", lines[i + 1]
                    ):
                        if len(new_lines) > 0 and new_lines[-1].strip() != "":
                            new_lines.append("\n")
                        while i < len(lines) and (
                            "|" in lines[i] and not lines[i].strip().startswith("```")
                        ):
                            new_lines.append(lines[i])
                            i += 1
                        if i < len(lines) and lines[i].strip() != "":
                            new_lines.append("\n")
                        continue
                new_lines.append(line)
                i += 1
            return new_lines

        def fix_markdown_headings(lines: list[str]) -> list[str]:
            """
            Ensures that headings have blank lines around them.

            Args:
                lines (list[str]): Lines of the Markdown file.

            Returns:
                list[str]: Modified lines.
            """

            new_lines = []
            for i, line in enumerate(lines):
                if re.match(r"^\s*#+ ", line):
                    if len(new_lines) > 0 and new_lines[-1].strip() != "":
                        new_lines.append("\n")
                    new_lines.append(line)
                    if i + 1 < len(lines) and lines[i + 1].strip() != "":
                        new_lines.append("\n")
                    continue
                new_lines.append(line)
            return new_lines

        def fix_markdown_lists(lines: list[str]) -> list[str]:
            """
            Ensures that lists have blank lines around them and consecutive lists do not merge with paragraphs.

            Args:
                lines (list[str]): Lines of the Markdown file.

            Returns:
                list[str]: Modified lines.
            """

            new_lines = []
            for i, line in enumerate(lines):
                if re.match(r"^\s*([-*+] |\d+\.)", line):
                    if len(new_lines) > 0 and new_lines[-1].strip() != "":
                        new_lines.append("\n")
                    new_lines.append(line)
                    if (
                        i + 1 < len(lines)
                        and not re.match(r"^\s*([-*+] |\d+\.)", lines[i + 1])
                        and lines[i + 1].strip() != ""
                    ):
                        new_lines.append("\n")
                    continue
                new_lines.append(line)
            return new_lines

        def normalize_markdown_whitespace(lines: list[str]) -> list[str]:
            """
            Removes excessive blank lines (max. 1 consecutive blank line).

            Args:
                lines (list[str]): Lines of the Markdown file.

            Returns:
                list[str]: Modified lines.
            """
            new_lines = []
            blank = False
            for line in lines:
                if line.strip() == "":
                    if not blank:
                        new_lines.append("\n")
                        blank = True
                else:
                    new_lines.append(line)
                    blank = False
            return new_lines

        lines = FileService.read_lines(self.markdown_file)

        lines = fix_markdown_tables(lines)
        lines = fix_markdown_headings(lines)
        lines = fix_markdown_lists(lines)
        lines = normalize_markdown_whitespace(lines)

        FileService.write_text("".join(lines), self.markdown_file)

    def replace_image_placeholders_with_markdown(self) -> None:
        """
        Replaces [IMAGE: ...] placeholders in the Markdown file with ![IMAGE](relative/path) format.

        Returns:
            None
        """

        content = FileService.read_text(self.markdown_file)

        def repl(match):
            img_name = match.group(1)
            rel_path = os.path.join(self.images_dir, img_name)
            if not rel_path:
                raise ValueError(f"No relative path found for image: {img_name}")
            return f"![IMAGE]({rel_path})"

        new_content = re.sub(r"\[IMAGE: ([^\]]+)\]", repl, content)

        FileService.write_text(new_content, self.markdown_file)


class UnstructuredConverter:
    """
    Converts Markdown files to JSON format for downstream processing.

    This class uses the Unstructured library to convert Markdown documents to structured JSON, supporting multiple processing strategies for optimal extraction. Designed for integration into document enrichment pipelines.

    Args:
        source_file (Path): Path to the source document file.
        markdown_file (Path): Path to the Markdown file to be processed.
        json_file (Path): Path to the output JSON file.
    """

    def __init__(self, source_file: Path, markdown_file: Path, json_file: Path) -> None:
        self.source_file = source_file
        self.markdown_file = markdown_file
        self.json_file = json_file

    def unstructured_process_markdown(self, strategy: str = "auto") -> None:
        """
        Converts a Markdown file to JSON format using the Unstructured library.

        Available options:
            - auto (default): Automatically selects the processing strategy based on document characteristics.
            - fast: Rule-based, fast text extraction; not recommended for image-heavy files.
            - hi_res: Model-based, considers document layout; recommended for more accurate element detection.
            - ocr_only: Model-based, uses only optical character recognition; mainly for image-based files.

        Args:
            strategy (str): Processing strategy (default: 'auto').

        Raises:
            FileNotFoundError: If the Markdown file is not found.
        """

        if not self.markdown_file.exists():
            raise FileNotFoundError(f"Markdown file not found: {self.markdown_file}")

        elements = partition(filename=self.markdown_file, strategy=strategy)

        FileService.write_json([el.to_dict() for el in elements], self.json_file)


class Enrichment:
    """
    Enriches JSON content with table and image summaries using external AI providers.

    This class integrates with the OpenAI API to generate structured summaries for tables and images found in documents. It supports robust error handling and parallel processing for efficient enrichment. Designed for flexible integration into document processing pipelines.

    Args:
        source_file (Path): Path to the source document file.
        json_file (Path): Path to the input JSON file to be enriched.
        client (OpenAI): OpenAI API client instance.
        enriched_json_file (Path): Path to the output enriched JSON file.
    """

    def __init__(
        self,
        source_file: Path,
        json_file: Path,
        client: OpenAI,
        enriched_json_file: Path,
    ) -> None:
        self.source_file = source_file
        self.json_file = json_file
        self.client = client
        self.enriched_json_file = enriched_json_file

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=10),
    )
    def call_openai_api(self, text_model: str, prompt: str, system_message: str) -> str:
        """
        Calls the OpenAI API with the specified prompt and system message.

        Args:
            text_model (str): The name of the OpenAI model.
            prompt (str): The user prompt.
            system_message (str): The system message.

        Returns:
            str: The response from the OpenAI API.

        Raises:
            OpenAIError: If the API call fails.
        """

        try:
            response = self.client.chat.completions.create(
                model=text_model,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt},
                ],
            )
            return response.choices[0].message.content
        except OpenAIError as e:
            raise OpenAIError(f"OpenAI API error: {str(e)}")

    def call_openai_vision_api(
        self, image_model: str, base64_image: str, detail: str
    ) -> str:
        """
        Calls the OpenAI Vision API to analyze an image.

        Args:
            image_model (str): The name of the OpenAI Vision model.
            base64_image (str): The image in Base64 encoded format.
            detail (str): The level of detail for the analysis. Possible values:
                - "low": fast, less detailed analysis
                - "high": detailed, thorough analysis
                - "auto": automatic selection of detail level

        Returns:
            str: The response from the OpenAI Vision API.
        """
        max_attempts = 5
        for attempt in range(1, max_attempts + 1):
            try:
                completion = self.client.chat.completions.create(
                    model=image_model,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": VISION_API_PROMPT,
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/jpeg;base64,{base64_image}",
                                        "detail": detail,
                                    },
                                },
                            ],
                        }
                    ],
                )
                return completion.choices[0].message.content
            except Exception as e:
                error_str = str(e)
                if (
                    "429" in error_str or "rate_limit_exceeded" in error_str
                ) and attempt < max_attempts:
                    wait_time = 10
                    time.sleep(wait_time)
                    continue
                return f"[Error during OpenAI Vision API call: {e}]"

    def summarize_table(self, text_model: str, html_content: str) -> str:
        """
        Summarizes the content of a table using the OpenAI API.

        Args:
            text_model (str): The name of the OpenAI text model.
            html_content (str): The HTML content of the table.

        Returns:
            str: The textual summary of the table.
        """
        soup = BeautifulSoup(html_content, "html.parser")

        prompt = TABLE_SUMMARY_PROMPT.format(table_data=soup)

        summary = self.call_openai_api(
            text_model=text_model,
            prompt=prompt,
            system_message=TABLE_SYSTEM_MESSAGE,
        )
        return summary

    def summarize_image(self, image_model: str, rel_image_path: str) -> str:
        """
        Generates a textual description of an image using the OpenAI Vision API.

        Args:
            image_model (str): The name of the OpenAI Vision model.
            rel_image_path (str): The relative path to the image.

        Returns:
            str: The textual description of the image.
        """
        abs_image_path = os.path.abspath(rel_image_path)
        base64_image = FileService.encode_image_base64(abs_image_path)

        return self.call_openai_vision_api(
            image_model=image_model, base64_image=base64_image, detail="low"
        )

    def enrich_json(
        self,
        process_table_with_ai: bool,
        process_images_with_ai: bool,
        text_model: str,
        image_model: str,
    ) -> None:
        """
        Enriches JSON content with table and image summaries.

        Args:
            process_table_with_ai (bool): Enable AI-based processing of tables.
            process_images_with_ai (bool): Enable AI-based processing of images.
        """
        try:
            elements = FileService.read_json(self.json_file)
        except FileNotFoundError:
            raise FileNotFoundError("JSON file not found")
        except json.JSONDecodeError:
            raise json.JSONDecodeError("Invalid JSON format")

        def process_element(element: Dict[str, Any]) -> Dict[str, Any]:
            """
            Processes a JSON element, which may contain a table or an image.

            If the element type is "Table" and contains a "text_as_html" field, a table summary is added.
            If the element type is "Image" and contains an "image_url" field, an image description is added.

            Args:
                element (Dict[str, Any]): The JSON element to process.

            Returns:
                Dict[str, Any]: The processed JSON element.

            Raises:
                Exception: If an error occurs during element processing.
            """
            try:
                if (
                    process_table_with_ai
                    and element["type"] == "Table"
                    and "text_as_html" in element["metadata"]
                ):
                    element["metadata"]["table_summary"] = self.summarize_table(
                        text_model=text_model,
                        html_content=element["metadata"]["text_as_html"],
                    )
                elif (
                    process_images_with_ai
                    and element["type"] == "Image"
                    and "image_url" in element["metadata"]
                ):
                    element["metadata"]["image_description"] = self.summarize_image(
                        image_model=image_model,
                        rel_image_path=element["metadata"]["image_url"],
                    )
            except Exception as e:
                raise RuntimeError(f"Error during element processing: {str(e)}")
            return element

        with ThreadPoolExecutor(max_workers=3) as executor:
            elements = list(executor.map(process_element, elements))

        FileService.write_json(elements, self.enriched_json_file)


class Export:
    """
    Exports enriched JSON content to plain text files for downstream processing.

    This class processes enriched JSON files, extracting table summaries, image descriptions, and text content, and writes them to TXT files in a format suitable for further use in the pipeline.

    Args:
        enriched_json_file (str): Path to the enriched JSON file.
        txt_file (str): Path to the output TXT file.
    """

    def __init__(self, enriched_json_file: str, txt_file: str) -> None:
        self.enriched_json_file = enriched_json_file
        self.txt_file = txt_file

    def export_text_from_enriched_json(self) -> None:
        """
        Exports text from an enriched JSON file to a TXT file.

        Processes the elements of the JSON file and writes the "table_summary" or "image_description"
        fields, as well as the "text" field (if no other content is present), to TXT format.

        Raises:
            FileNotFoundError: If the JSON file is not found.
        """
        try:
            elements = FileService.read_json(self.enriched_json_file)
        except FileNotFoundError:
            raise FileNotFoundError(f"JSON file not found: {self.enriched_json_file}")

        paragraphs = []
        for element in elements:
            has_metadata_content = False
            if element.get("metadata"):
                if "table_summary" in element["metadata"]:
                    paragraphs.append(element["metadata"]["table_summary"].strip())
                    has_metadata_content = True
                if "image_description" in element["metadata"]:
                    paragraphs.append(element["metadata"]["image_description"].strip())
                    has_metadata_content = True
            if element.get("text") and not has_metadata_content:
                text = element["text"].strip()
                if text:
                    paragraphs.append(text)

        output_text = "\n\n".join(paragraphs)
        FileService.write_text(output_text, self.txt_file)
