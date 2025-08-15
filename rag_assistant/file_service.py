import base64
import json
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.oxml.ns import qn


class FileService:
    """
    Utility class for file operations with static methods.

    Provides methods for reading, writing, and processing files, including DOCX image extraction and encoding images to base64.
    """

    @staticmethod
    def extract_images_from_docx(source_file: Path, images_dir: Path) -> None:
        """
        Extracts images from a DOCX file to the specified directory and inserts placeholders.

        Args:
            source_file (Path): Path to the source DOCX file.
            images_dir (Path): Path to the target images directory.

        Returns:
            None
        """

        doc = Document(str(source_file))
        images_dir.mkdir(parents=True, exist_ok=True)

        rels = doc.part.rels
        image_map = {}
        parent_map = {}

        for para in doc.paragraphs:
            for run in para.runs:
                if "graphic" in run._element.xml:
                    drawing_elems = run._element.findall(
                        ".//w:drawing", namespaces=run._element.nsmap
                    )
                    for drawing in drawing_elems:
                        blip = drawing.find(
                            ".//a:blip",
                            namespaces={
                                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                            },
                        )
                        if blip is not None:
                            rId = blip.get(qn("r:embed"))
                            if rId:
                                parent_map[rId] = run
        image_count = 0
        for rel_id, rel in rels.items():
            if (
                rel.reltype
                == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
            ):
                image_count += 1
                img_data = rel.target_part.blob
                img_ext = rel.target_part.content_type.split("/")[-1]
                img_name = f"image_{image_count}.{img_ext}"
                img_path = images_dir / img_name
                FileService.write_binary(img_data, img_path)
                image_map[rel_id] = img_name

                if rel_id in parent_map:
                    parent_run = parent_map[rel_id]
                    para = parent_run._parent
                    para.add_run(f" [IMAGE: {img_name}]")

        doc.save(str(source_file))

    @staticmethod
    def write_json(data: List[Dict], output_path: Path) -> None:
        """
        Writes JSON data to a file.

        Args:
            data (List[Dict]): The JSON data to write.
            output_path (Path): Path to the output file.

        Returns:
            None

        Raises:
            RuntimeError: If an error occurs while writing the file.
        """
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except (OSError, TypeError) as e:
            raise RuntimeError(
                f"Hiba a JSON fájl írása közben: {e} (output_path={output_path})"
            )

    @staticmethod
    def write_binary(data: bytes, output_path: Path) -> None:
        """
        Writes binary data to a file.

        Args:
            data (bytes): The binary data to write.
            output_path (Path): Path to the output file.

        Returns:
            None
        """
        with open(output_path, "wb") as f:
            f.write(data)

    @staticmethod
    def write_text(text: str, output_path: Path) -> None:
        """
        Writes text to a file.

        Args:
            text (str): The text to write.
            output_path (Path): Path to the output file.

        Returns:
            None
        """
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

    @staticmethod
    def read_text(input_path: Path) -> str:
        """
        Reads text from a file.

        Args:
            input_path (Path): Path to the input file.

        Returns:
            str: The text read from the file.
        """
        with open(input_path, "r", encoding="utf-8") as f:
            return f.read()

    @staticmethod
    def read_lines(input_path: Path) -> List[str]:
        """
        Reads lines from a file.

        Args:
            input_path (Path): Path to the input file.

        Returns:
            List[str]: List of lines read from the file.
        """
        with open(input_path, "r", encoding="utf-8") as f:
            return f.readlines()

    @staticmethod
    def read_json(input_path: Path) -> List[Dict]:
        """
        Reads JSON data from a file.

        Args:
            input_path (Path): Path to the input file.

        Returns:
            List[Dict]: JSON data read from the file.
        """
        with open(input_path, "r", encoding="utf-8") as f:
            return json.load(f)

    @staticmethod
    def encode_image_base64(image_path: Path) -> str:
        """
        Encodes an image file to base64 string.

        Args:
            image_path (Path): Path to the image file.

        Returns:
            str: Base64-encoded image string.
        """
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode("utf-8")
